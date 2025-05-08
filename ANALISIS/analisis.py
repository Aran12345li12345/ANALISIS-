import os
from flask import Flask, request, render_template, redirect, url_for
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def analyze_pdf(file_path):
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return f"PDF Analysis: {len(reader.pages)} pages, {len(text.split())} words."
    except Exception as e:
        return f"Error analyzing PDF: {e}"

def analyze_docx(file_path):
    try:
        doc = Document(file_path)
        text = " ".join([paragraph.text for paragraph in doc.paragraphs])
        return f"Word Document Analysis: {len(doc.paragraphs)} paragraphs, {len(text.split())} words."
    except Exception as e:
        return f"Error analyzing Word document: {e}"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return redirect(request.url)
    file = request.files['file']
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        if filename.endswith('.pdf'):
            analysis_result = analyze_pdf(file_path)
        elif filename.endswith('.docx'):
            analysis_result = analyze_docx(file_path)
        else:
            analysis_result = "Unsupported file type."

        return render_template('index.html', result=analysis_result)
    else:
        return redirect(request.url)

if __name__ == '__main__':
    app.run(debug=True)